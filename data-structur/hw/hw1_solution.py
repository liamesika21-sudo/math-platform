from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(12)
font.color.rgb = RGBColor(0, 0, 0)
# For Hebrew/RTL support
rPr = style.element.get_or_add_rPr()
rFonts = rPr.makeelement(qn('w:rFonts'), {qn('w:cs'): 'Arial', qn('w:ascii'): 'Arial', qn('w:hAnsi'): 'Arial'})
rPr.append(rFonts)

def add_rtl_paragraph(text, bold=False, size=12, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=6):
    p = doc.add_paragraph()
    p.alignment = alignment
    # Set RTL
    pPr = p._element.get_or_add_pPr()
    bidi = pPr.makeelement(qn('w:bidi'), {})
    pPr.append(bidi)

    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = bold
    # CS font for Hebrew
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.makeelement(qn('w:rFonts'), {qn('w:cs'): 'Arial'})
    rPr.append(rFonts)
    # Set RTL on run
    rtl = rPr.makeelement(qn('w:rtl'), {})
    rPr.append(rtl)

    from docx.shared import Pt as PtShared
    p.paragraph_format.space_after = PtShared(space_after)
    return p

def add_mixed_paragraph(parts, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=6):
    """parts is a list of (text, bold, is_rtl, size) tuples"""
    p = doc.add_paragraph()
    p.alignment = alignment
    pPr = p._element.get_or_add_pPr()
    bidi = pPr.makeelement(qn('w:bidi'), {})
    pPr.append(bidi)

    for text, bold, is_rtl, size in parts:
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.bold = bold
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.makeelement(qn('w:rFonts'), {qn('w:cs'): 'Arial'})
        rPr.append(rFonts)
        if is_rtl:
            rtl = rPr.makeelement(qn('w:rtl'), {})
            rPr.append(rtl)

    from docx.shared import Pt as PtShared
    p.paragraph_format.space_after = PtShared(space_after)
    return p

def add_ltr_paragraph(text, bold=False, size=12, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=6):
    p = doc.add_paragraph()
    p.alignment = alignment
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.makeelement(qn('w:rFonts'), {qn('w:cs'): 'Arial'})
    rPr.append(rFonts)
    from docx.shared import Pt as PtShared
    p.paragraph_format.space_after = PtShared(space_after)
    return p

# ==================== HEADER ====================
add_rtl_paragraph("מבני נתונים - מטלה 1", bold=True, size=16, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_rtl_paragraph("אוניברסיטת רייכמן, אביב 2026", bold=False, size=13, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_rtl_paragraph("", size=6)

# ==================== QUESTION A ====================
add_rtl_paragraph("שאלה A: יסודות מתמטיים (20 נקודות)", bold=True, size=14)

# A1
add_rtl_paragraph("סעיף 1:", bold=True)
add_rtl_paragraph("נתונה המשוואה:")
add_ltr_paragraph("log₂(x) + log₄(x) + log₈(x) = 11", size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_rtl_paragraph("פתרון:", bold=True)
add_rtl_paragraph("נמיר את כל הלוגריתמים לבסיס 2 באמצעות נוסחת המרת בסיס:")
add_ltr_paragraph("log₄(x) = log₂(x) / log₂(4) = log₂(x) / 2", alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_ltr_paragraph("log₈(x) = log₂(x) / log₂(8) = log₂(x) / 3", alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_rtl_paragraph("נציב במשוואה:")
add_ltr_paragraph("log₂(x) + log₂(x)/2 + log₂(x)/3 = 11", alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_rtl_paragraph("נוציא גורם משותף:")
add_ltr_paragraph("log₂(x) · (1 + 1/2 + 1/3) = 11", alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_ltr_paragraph("log₂(x) · (6/6 + 3/6 + 2/6) = 11", alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_ltr_paragraph("log₂(x) · (11/6) = 11", alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_ltr_paragraph("log₂(x) = 6", alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_ltr_paragraph("x = 2⁶ = 64", alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)

add_rtl_paragraph("")

# A2
add_rtl_paragraph("סעיף 2:", bold=True)
add_rtl_paragraph("נתון ש-n הוא חזקה של 2 (כלומר n = 2^k עבור k ≥ 1 שלם). נחשב את הסכום:")
add_ltr_paragraph("Σ (i=0 to log₂n - 1) n/2ⁱ", alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_rtl_paragraph("פתרון:", bold=True)
add_rtl_paragraph("נוציא את n כגורם משותף:")
add_ltr_paragraph("= n · Σ (i=0 to log₂n - 1) (1/2)ⁱ", alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_rtl_paragraph("זוהי טור הנדסי עם איבר ראשון a = 1, מנה q = 1/2, ומספר איברים log₂(n).")
add_rtl_paragraph("לפי נוסחת סכום טור הנדסי:")
add_ltr_paragraph("= n · (1 - (1/2)^(log₂n)) / (1 - 1/2)", alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_ltr_paragraph("= n · (1 - 1/n) / (1/2)", alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_rtl_paragraph("שימו לב: (1/2)^(log₂n) = 1/(2^(log₂n)) = 1/n")
add_ltr_paragraph("= n · 2 · (1 - 1/n)", alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_ltr_paragraph("= 2n - 2 = 2(n - 1)", alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)

add_rtl_paragraph("")

# A3
add_rtl_paragraph("סעיף 3:", bold=True)
add_rtl_paragraph("נפשט את הביטוי:")
add_ltr_paragraph("2^(log₄(x))", alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_rtl_paragraph("פתרון:", bold=True)
add_rtl_paragraph("נמיר את הלוגריתם לבסיס 2:")
add_ltr_paragraph("log₄(x) = log₂(x) / log₂(4) = log₂(x) / 2", alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_rtl_paragraph("נציב:")
add_ltr_paragraph("2^(log₄(x)) = 2^(log₂(x) / 2)", alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_ltr_paragraph("= (2^(log₂(x)))^(1/2)", alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_ltr_paragraph("= x^(1/2)", alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_ltr_paragraph("= √x", alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)

add_rtl_paragraph("")

# ==================== QUESTION B ====================
add_rtl_paragraph("שאלה B: הוכחה באינדוקציה - חלק 1 (30 נקודות)", bold=True, size=14)

# B1
add_rtl_paragraph("סעיף 1:", bold=True)
add_rtl_paragraph("נוכיח באינדוקציה על d שמספר הנדבקים החדשים ביום d הוא לכל היותר 2^(d-1).")
add_rtl_paragraph("")

add_rtl_paragraph("בסיס האינדוקציה (d = 1):", bold=True)
add_rtl_paragraph("ביום d = 1, בדיוק אדם אחד נדבק (חולה אפס).")
add_rtl_paragraph("נבדוק: 2^(1-1) = 2⁰ = 1. אכן, 1 ≤ 1. הבסיס מתקיים. ✓")
add_rtl_paragraph("")

add_rtl_paragraph("הנחת האינדוקציה:", bold=True)
add_rtl_paragraph("נניח שעבור יום d כלשהו, מספר הנדבקים החדשים הוא לכל היותר 2^(d-1).")
add_rtl_paragraph("")

add_rtl_paragraph("צעד האינדוקציה (d → d+1):", bold=True)
add_rtl_paragraph("נוכיח שמספר הנדבקים החדשים ביום d+1 הוא לכל היותר 2^d.")
add_rtl_paragraph("לפי הנתון, כל אדם שנדבק ביום d מדביק לכל היותר 2 אנשים חדשים ביום d+1.")
add_rtl_paragraph("לפי הנחת האינדוקציה, מספר הנדבקים ביום d הוא לכל היותר 2^(d-1).")
add_rtl_paragraph("לכן, מספר הנדבקים החדשים ביום d+1 הוא לכל היותר:")
add_ltr_paragraph("2^(d-1) · 2 = 2^d = 2^((d+1)-1)", alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_rtl_paragraph("וזה בדיוק מה שרצינו להוכיח. ∎")
add_rtl_paragraph("")

# B2a
add_rtl_paragraph("סעיף 2:", bold=True)

add_rtl_paragraph("(א) טענה 1: ההוכחה שגויה.", bold=True)
add_rtl_paragraph("הבעיה היא במעבר מבסיס האינדוקציה (n=2) לצעד האינדוקציה (n=3). כאשר n=2, "
                   "צעד האינדוקציה מנסה להוכיח עבור n+1=3 נקודות. עם 3 נקודות p₁, p₂, p₃, "
                   "נפעיל את הנחת האינדוקציה על {p₁, p₂} ועל {p₂, p₃}. "
                   "אך הנקודה המשותפת היחידה בין שתי הקבוצות היא p₂ בלבד. "
                   "נקודה משותפת אחת אינה מספיקה כדי להסיק ש-ℓ₁ = ℓ₂, "
                   "שכן שני ישרים שונים יכולים לחלוק נקודה אחת. "
                   "הטיעון דורש לפחות 2 נקודות משותפות, מה שמתקיים רק כאשר n ≥ 3.")
add_rtl_paragraph("")

# B2b
add_rtl_paragraph("(ב) טענה 2: ההוכחה נכונה.", bold=True)
add_rtl_paragraph("כל השלבים בהוכחה תקינים:")
add_rtl_paragraph("- בסיס האינדוקציה: T(2) = 4 ≤ 10·2·log₂(2) = 20. ✓")
add_rtl_paragraph("- הנחת האינדוקציה מוגדרת כראוי עבור n/2.")
add_rtl_paragraph("- צעד האינדוקציה: השתמשו נכון בחסם הרקורסיבי ובהנחת האינדוקציה, "
                   "והאלגברה בשלב (iv) נכונה:")
add_ltr_paragraph("T(n) ≤ 10n(log₂(n/2) + 1) = 10n(log₂n - 1 + 1) = 10n·log₂n  ✓", alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_rtl_paragraph("")

# ==================== QUESTION C ====================
add_rtl_paragraph("שאלה C: הוכחה באינדוקציה - חלק 2 (30 נקודות)", bold=True, size=14)

# C1
add_rtl_paragraph("סעיף 1: הוכיחו באינדוקציה שלכל n ≥ 4 מתקיים n! > 2ⁿ.", bold=True)
add_rtl_paragraph("")

add_rtl_paragraph("בסיס האינדוקציה (n = 4):", bold=True)
add_rtl_paragraph("4! = 24, ו-2⁴ = 16. אכן 24 > 16. ✓")
add_rtl_paragraph("")

add_rtl_paragraph("הנחת האינדוקציה:", bold=True)
add_rtl_paragraph("נניח שעבור k ≥ 4 כלשהו מתקיים k! > 2^k.")
add_rtl_paragraph("")

add_rtl_paragraph("צעד האינדוקציה (k → k+1):", bold=True)
add_rtl_paragraph("נוכיח ש-(k+1)! > 2^(k+1).")
add_ltr_paragraph("(k+1)! = (k+1) · k!", alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_rtl_paragraph("לפי הנחת האינדוקציה, k! > 2^k, ולכן:")
add_ltr_paragraph("(k+1)! = (k+1) · k! > (k+1) · 2^k", alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_rtl_paragraph("מכיוון ש-k ≥ 4, מתקיים k+1 ≥ 5 > 2, ולכן:")
add_ltr_paragraph("(k+1) · 2^k > 2 · 2^k = 2^(k+1)", alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_rtl_paragraph("קיבלנו (k+1)! > 2^(k+1), כנדרש. ∎")
add_rtl_paragraph("")

# C2
add_rtl_paragraph("סעיף 2: הוכיחו באינדוקציה שלכל n חיובי:", bold=True)
add_ltr_paragraph("Σ (i=2 to n) i·2ⁱ = 2^(n+1)·(n-1)", alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_rtl_paragraph("")

add_rtl_paragraph("בסיס האינדוקציה (n = 2):", bold=True)
add_rtl_paragraph("אגף שמאל: 2·2² = 2·4 = 8.")
add_rtl_paragraph("אגף ימין: 2^(2+1)·(2-1) = 2³·1 = 8.")
add_rtl_paragraph("אכן 8 = 8. ✓")
add_rtl_paragraph("")

add_rtl_paragraph("הנחת האינדוקציה:", bold=True)
add_rtl_paragraph("נניח שעבור k ≥ 2 כלשהו מתקיים:")
add_ltr_paragraph("Σ (i=2 to k) i·2ⁱ = 2^(k+1)·(k-1)", alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_rtl_paragraph("")

add_rtl_paragraph("צעד האינדוקציה (k → k+1):", bold=True)
add_rtl_paragraph("נוכיח שמתקיים:")
add_ltr_paragraph("Σ (i=2 to k+1) i·2ⁱ = 2^(k+2)·k", alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_rtl_paragraph("נפרק את הסכום:")
add_ltr_paragraph("Σ (i=2 to k+1) i·2ⁱ = Σ (i=2 to k) i·2ⁱ + (k+1)·2^(k+1)", alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_rtl_paragraph("נציב את הנחת האינדוקציה:")
add_ltr_paragraph("= 2^(k+1)·(k-1) + (k+1)·2^(k+1)", alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_rtl_paragraph("נוציא גורם משותף 2^(k+1):")
add_ltr_paragraph("= 2^(k+1) · [(k-1) + (k+1)]", alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_ltr_paragraph("= 2^(k+1) · 2k", alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_ltr_paragraph("= 2^(k+2) · k", alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_rtl_paragraph("וזהו בדיוק 2^((k+1)+1)·((k+1)-1), כנדרש. ∎")
add_rtl_paragraph("")

# C3
add_rtl_paragraph("סעיף 3: הוכחת נכונות אלגוריתם חיפוש בינארי", bold=True)
add_rtl_paragraph("")

add_rtl_paragraph("טענה:", bold=True)
add_rtl_paragraph("בהינתן ערך x ומערך ממוין A, הקריאה BSearch(x, A, i, j) מחזירה אינדקס i ≤ m ≤ j "
                   "כך ש-A[m] = x, או מחזירה -1 אם x אינו במערך. נוכיח באינדוקציה על גודל "
                   "התת-מערך n = j - i + 1.")
add_rtl_paragraph("")

add_rtl_paragraph("בסיס האינדוקציה (n ≤ 0, כלומר j - i < 0):", bold=True)
add_rtl_paragraph("כאשר j - i < 0, התת-מערך ריק ולכן x אינו בו. "
                   "בשורה 2, התנאי j - i < 0 מתקיים, ולכן האלגוריתם מחזיר -1 בשורה 3. "
                   "זהו ערך נכון כיוון שהתת-מערך ריק. ✓")
add_rtl_paragraph("")

add_rtl_paragraph("הנחת האינדוקציה:", bold=True)
add_rtl_paragraph("נניח שהאלגוריתם BSearch עובד נכון לכל תת-מערך בגודל קטן מ-n.")
add_rtl_paragraph("")

add_rtl_paragraph("צעד האינדוקציה:", bold=True)
add_rtl_paragraph("נוכיח שהאלגוריתם עובד נכון לתת-מערך בגודל n = j - i + 1 ≥ 1.")
add_rtl_paragraph("")
add_rtl_paragraph("מכיוון ש-n ≥ 1, מתקיים j - i ≥ 0, ולכן התנאי בשורה 2 אינו מתקיים.")
add_rtl_paragraph("בשורה 4, האלגוריתם מחשב m = ⌊(i+j)/2⌋. שימו לב ש-i ≤ m ≤ j.")
add_rtl_paragraph("")

add_rtl_paragraph("ישנם שלושה מקרים:")
add_rtl_paragraph("")

add_rtl_paragraph("מקרה 1: x = A[m] (שורה 5).", bold=True)
add_rtl_paragraph("האלגוריתם מחזיר m בשורה 6. מכיוון ש-A[m] = x ו-i ≤ m ≤ j, זהו ערך נכון. ✓")
add_rtl_paragraph("")

add_rtl_paragraph("מקרה 2: x < A[m] (שורה 7).", bold=True)
add_rtl_paragraph("מכיוון שהמערך A ממוין, אם x קיים במערך הוא חייב להיות בתת-המערך A[i..m-1]. "
                   "האלגוריתם קורא ל-BSearch(x, A, i, m-1) בשורה 8.")
add_rtl_paragraph("גודל התת-מערך החדש הוא (m-1) - i + 1 = m - i.")
add_rtl_paragraph("מכיוון ש-m = ⌊(i+j)/2⌋ ≤ j, מתקיים m - i ≤ j - i < j - i + 1 = n.")
add_rtl_paragraph("לכן גודל התת-מערך החדש קטן מ-n, ולפי הנחת האינדוקציה הקריאה הרקורסיבית מחזירה תוצאה נכונה. ✓")
add_rtl_paragraph("")

add_rtl_paragraph("מקרה 3: x > A[m] (שורות 9-10).", bold=True)
add_rtl_paragraph("מכיוון שהמערך A ממוין, אם x קיים במערך הוא חייב להיות בתת-המערך A[m+1..j]. "
                   "האלגוריתם קורא ל-BSearch(x, A, m+1, j) בשורה 10.")
add_rtl_paragraph("גודל התת-מערך החדש הוא j - (m+1) + 1 = j - m.")
add_rtl_paragraph("מכיוון ש-m = ⌊(i+j)/2⌋ ≥ i, מתקיים j - m ≤ j - i < j - i + 1 = n.")
add_rtl_paragraph("לכן גודל התת-מערך החדש קטן מ-n, ולפי הנחת האינדוקציה הקריאה הרקורסיבית מחזירה תוצאה נכונה. ✓")
add_rtl_paragraph("")

add_rtl_paragraph("בכל שלושת המקרים האלגוריתם מחזיר תוצאה נכונה, וזה משלים את ההוכחה. ∎")
add_rtl_paragraph("")

# ==================== QUESTION D ====================
add_rtl_paragraph("שאלה D: מבני נתונים (20 נקודות)", bold=True, size=14)

# D1
add_rtl_paragraph("סעיף 1: פסאודו-קוד להיפוך רשימה מקושרת חד-כיוונית", bold=True)
add_rtl_paragraph("")

add_ltr_paragraph("""function ReverseList(L):
    prev ← null
    curr ← L.head
    while curr ≠ null do
        next ← curr.next
        curr.next ← prev
        prev ← curr
        curr ← next
    L.head ← prev""", size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT)

add_rtl_paragraph("")
add_rtl_paragraph("הסבר:", bold=True)
add_rtl_paragraph("האלגוריתם עובר על הרשימה מההתחלה לסוף. בכל צעד, הוא שומר מצביע לצומת הבא (next), "
                   "הופך את כיוון המצביע של הצומת הנוכחי כך שיצביע לצומת הקודם (prev), "
                   "ואז מתקדם הלאה. בסיום, head מצביע לאיבר האחרון של הרשימה המקורית, "
                   "שהוא כעת האיבר הראשון ברשימה ההפוכה.")
add_rtl_paragraph("")

# D2
add_rtl_paragraph("סעיף 2: ניתוח סיבוכיות", bold=True)
add_rtl_paragraph("")

add_rtl_paragraph("סיבוכיות זמן: O(n)", bold=True)
add_rtl_paragraph("האלגוריתם מבצע מעבר יחיד על כל הרשימה. לולאת ה-while רצה בדיוק n פעמים "
                   "(פעם אחת לכל צומת ברשימה). בכל איטרציה מתבצע מספר קבוע של פעולות "
                   "(השמות מצביעים). לכן סיבוכיות הזמן היא O(n).")
add_rtl_paragraph("")

add_rtl_paragraph("סיבוכיות מקום: O(1)", bold=True)
add_rtl_paragraph("האלגוריתם משתמש רק בשלושה משתני עזר: prev, curr, ו-next. "
                   "מספר המשתנים קבוע ואינו תלוי בגודל הקלט n. "
                   "לא נעשה שימוש במערך, מבנה נתונים נוסף, או רקורסיה. "
                   "לכן סיבוכיות המקום הנוספת היא O(1).")

# Save
output_path = "/Users/liamesika/Desktop/infi/data-structur/hw/hw1_solution.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")
